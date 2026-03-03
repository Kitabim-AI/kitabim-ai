"""Service for extracting text and cover from .docx files."""
from __future__ import annotations

import logging
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from app.utils.text import normalize_uyghur_chars

logger = logging.getLogger(__name__)

_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
# Types that are always present in footnotes.xml but carry no real content
_FOOTNOTE_SKIP_TYPES = {"separator", "continuationSeparator", "continuationNotice"}


def _load_footnotes(docx_path: Path) -> dict[str, str]:
    """
    Parse word/footnotes.xml and return a mapping of {id: text}.
    Skips the built-in separator footnotes (id -1 and 0).
    Returns an empty dict if the file doesn't exist.
    """
    result: dict[str, str] = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            if "word/footnotes.xml" not in z.namelist():
                return result
            xml_bytes = z.read("word/footnotes.xml")
    except Exception as e:
        logger.warning("docx_service: could not read footnotes path=%s error=%s", docx_path, e)
        return result

    root = ET.fromstring(xml_bytes.decode("utf-8"))
    for fn in root.findall(f"{{{_NS}}}footnote"):
        ftype = fn.get(f"{{{_NS}}}type", "normal")
        if ftype in _FOOTNOTE_SKIP_TYPES:
            continue
        fid = fn.get(f"{{{_NS}}}id", "")
        text = "".join(t.text or "" for t in fn.findall(f".//{{{_NS}}}t")).strip()
        if fid and text:
            result[fid] = text
    return result


def _is_heading_para(para_el) -> bool:
    """
    Return True if the paragraph looks like a section heading.
    Heuristic: every run in the paragraph has bold enabled (w:b),
    and the total text is short enough to be a title (< 200 chars).
    """
    runs = para_el.findall(f"{{{_NS}}}r")
    if not runs:
        return False
    full_text = "".join(
        (t.text or "")
        for r in runs
        for t in r.findall(f"{{{_NS}}}t")
    )
    if not full_text.strip() or len(full_text.strip()) >= 200:
        return False
    return all(r.find(f".//{{{_NS}}}b") is not None for r in runs)


def extract_docx_pages(path: Path) -> list[str]:
    """
    Split a .docx file into pages using lastRenderedPageBreak markers.

    - Splits at run level to handle paragraphs spanning multiple pages.
    - Bold short paragraphs are prefixed with '## ' as section headings.
    - Footnote texts are appended at the bottom of the page they appear on.
    - Falls back to fixed 3000-char blocks if no page-break markers are found.
    """
    from docx import Document

    doc = Document(path)
    footnotes = _load_footnotes(path)
    logger.debug("docx_service: loaded %d footnotes from %s", len(footnotes), path.name)

    pages: list[str] = []
    current_page: list[str] = []
    current_footnotes: list[str] = []  # footnotes referenced on the current page

    def _flush_page():
        lines = list(current_page)
        if current_footnotes:
            lines.append("---")
            lines.extend(current_footnotes)
        pages.append("\n".join(lines))
        current_page.clear()
        current_footnotes.clear()

    for para in doc.paragraphs:
        is_heading = _is_heading_para(para._p)
        para_chunks: list[list[str]] = [[]]
        para_fn_ids: list[list[str]] = [[]]  # footnote ids collected per chunk

        for run in para._p.findall(f"{{{_NS}}}r"):
            lrpb = run.find(f"{{{_NS}}}lastRenderedPageBreak")
            t = run.find(f"{{{_NS}}}t")
            fn_ref = run.find(f"{{{_NS}}}footnoteReference")

            text = (t.text or "") if t is not None else ""
            fn_id = fn_ref.get(f"{{{_NS}}}id") if fn_ref is not None else None

            if lrpb is not None:
                para_chunks.append([text])
                para_fn_ids.append([fn_id] if fn_id else [])
            else:
                para_chunks[-1].append(text)
                if fn_id:
                    para_fn_ids[-1].append(fn_id)

        for idx, (chunk, fn_ids) in enumerate(zip(para_chunks, para_fn_ids)):
            chunk_text = "".join(chunk).strip()

            if idx > 0:
                _flush_page()

            # Collect footnotes for this chunk
            for fid in fn_ids:
                fn_text = footnotes.get(fid)
                if fn_text:
                    current_footnotes.append(fn_text)

            if not chunk_text:
                continue

            if is_heading:
                chunk_text = f"## {chunk_text}"

            current_page.append(chunk_text)

    if current_page or current_footnotes:
        _flush_page()

    # Fallback: no page break markers found, split by fixed block size
    if len(pages) <= 1:
        full_text = pages[0] if pages else ""
        block_size = 3000
        pages = [full_text[i:i + block_size] for i in range(0, len(full_text), block_size)]

    result = [normalize_uyghur_chars(p.strip()) for p in pages if p.strip()]
    logger.info("docx_service: extracted %d pages from %s", len(result), path.name)
    return result


def extract_docx_cover(docx_path: Path, cover_path: Path) -> bool:
    """
    Extract the first inline image from a .docx file and save it as cover_path.
    Returns True on success, False if no image found or on error.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                with open(cover_path, "wb") as f:
                    f.write(rel.target_part.blob)
                return True
    except Exception as e:
        logger.warning("docx_service: failed to extract cover path=%s error=%s", docx_path, e)
    return False
